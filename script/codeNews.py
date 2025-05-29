import os
from docx import Document
from openai import OpenAI
from tqdm import tqdm  
import pandas as pd
import time
import pandas as pd
import re
import numpy as np
import matplotlib.pyplot as plt
from collections import Counter


OPENAI_API_KEY='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'


def call_openai(prompt):
    """ Llama a la API de OpenAI con el modelo GPT-4o """
    client = OpenAI(api_key=OPENAI_API_KEY)

    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="gpt-4o"
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        return f"Error en OpenAI: {str(e)}"


def extraer_texto_docx(ruta_archivo):
    """ Extrae el texto de un archivo .docx """
    try:
        doc = Document(ruta_archivo)
        texto = "\n".join([p.text for p in doc.paragraphs])
        texto = re.sub(r'\d+ª Noticia\s*', '', texto)
        return texto.strip()
    except Exception as e:
        return f"Error al leer {ruta_archivo}: {str(e)}"


def extraer_numero(archivo):
    """ Extrae el número de un nombre de archivo como 'Noticia 1.docx' -> 1 """
    numeros = re.findall(r'\d+', archivo)  # Encuentra todos los números en el nombre
    return int(numeros[0]) if numeros else float('inf')  # Si no hay números, lo manda al final


def obtener_archivos(carpeta):
    """ Recorre todos los archivos .docx en la carpeta y extrae su texto con barra de progreso """
    if not os.path.exists(carpeta):
        print(f"Error: La carpeta '{carpeta}' no existe.")
        return

    archivos = [archivo for archivo in os.listdir(carpeta) if archivo.endswith(".docx")]
    archivos = sorted(archivos, key=extraer_numero)

    if not archivos:
        print("No se encontraron archivos .docx en la carpeta.")
        return
    
    return archivos


def obtener_noticias(carpeta, archivos):
    # Extraer noticias
    columns = ["Noticia", "Contenido", "Análisis Cuantitativo y Cualitativo", "Opinión de Agentes", 
                "Análisis de Sentimientos", "Análisis Sentimientos Cuantitativo", "Provincia"]
    contador_noticia = 1  
    noticias = pd.DataFrame(columns = columns)

    for archivo in tqdm(archivos, desc="Procesando archivos", unit="archivo"):
        ruta_completa = os.path.join(carpeta, archivo)
        contenido = extraer_texto_docx(ruta_completa)
        nuevo_indice = len(noticias)
        noticias.loc[nuevo_indice, "Noticia"] = f"Noticia {contador_noticia}"
        noticias.loc[nuevo_indice, "Contenido"] = contenido
        contador_noticia += 1

    return noticias


def analisis_pregunta(noticias, prompt, columna):
    pattern = re.compile(r"^(lo\s+siento|i['’]?m\s+sorry)", re.IGNORECASE)

    noticias_copy = noticias.copy()  # Copia el DataFrame para no modificar el original
    # Recorremos cada fila del DataFrame
    for index, row in tqdm(noticias.iterrows(), total=noticias.shape[0], desc="Procesando " + columna, unit="noticia"):
        contenido_noticia = row["Contenido"]       
        prompt_completo = prompt + contenido_noticia
        resultado = call_openai(prompt_completo)
        max_attempts = 20  # Número máximo de reintentos permitidos
        attempt = 0
        while re.match(pattern, resultado) and attempt < max_attempts:
            attempt += 1
            resultado = call_openai(prompt_completo)
        noticias_copy.loc[index, columna] = str(resultado) if resultado is not None else ""
        time.sleep(1)
    return noticias_copy

def obtener_palabras(texto):
    # Convertimos a minúsculas y usamos una expresión regular para extraer solo palabras
    return re.findall(r'\w+', texto.lower())

def porcentaje_similitud_longest(texto1, texto2):
    # Dividimos los textos en tokens
    tokens1 = texto1.split()
    tokens2 = texto2.split()
    # Calculamos la longitud de la subcadena contigua más larga
    longest = longest_common_substring(tokens1, tokens2)
    # Usamos el número de tokens del texto más extenso como denominador
    max_tokens = max(len(tokens1), len(tokens2))
    if max_tokens == 0:
        return 0
    porcentaje = longest / max_tokens * 100
    return porcentaje


def longest_common_substring(tokens1, tokens2):
    """
    Calcula la longitud de la subcadena contigua más larga entre dos listas de tokens.
    """
    m, n = len(tokens1), len(tokens2)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    longest = 0
    for i in range(m):
        for j in range(n):
            if tokens1[i] == tokens2[j]:
                dp[i+1][j+1] = dp[i][j] + 1
                if dp[i+1][j+1] > longest:
                    longest = dp[i+1][j+1]
            else:
                dp[i+1][j+1] = 0
    return longest


def similitud_subcadena(text1, text2):
    """
    Calcula la similitud en función de la longitud de la subcadena contigua más larga
    utilizando el texto original sin limpieza previa.
    """
    tokens1 = text1.split()
    tokens2 = text2.split()
    return longest_common_substring(tokens1, tokens2)

def analisis_pregunta_explicacion_sentimientos(noticias, prompt, columna):
    pattern = re.compile(r"^(lo\s+siento|i['’]?m\s+sorry)", re.IGNORECASE)

    noticias_copy = noticias.copy()  # Copia el DataFrame para no modificar el original
    # Recorremos cada fila del DataFrame
    for index, row in tqdm(noticias.iterrows(), total=noticias.shape[0], desc="Procesando " + columna, unit="noticia"):
        contenido_noticia = row["Contenido"]       
        clasificacion = f"Esta noticia tiene una clasificacion de {row["Sentimiento IA"]}. La noticia es la siguiente: " 
        prompt_completo = prompt + clasificacion + contenido_noticia
        resultado = call_openai(prompt_completo)
        max_attempts = 20  # Número máximo de reintentos permitidos
        attempt = 0
        while re.match(pattern, resultado) and attempt < max_attempts:
            attempt += 1
            resultado = call_openai(prompt_completo)
        noticias_copy.loc[index, columna] = str(resultado) if resultado is not None else ""
        time.sleep(1)
    return noticias_copy

def concatenar_salida(noticias):
    noticias_copy = noticias.copy()  # Copia el DataFrame para no modificar el original
    for index, row in noticias.iterrows():
        union = "\nLos agentes a analizar son: \n"
        noticias_copy.loc[index, 'Contenido'] = str(noticias_copy.loc[index, 'Contenido']) + union + str(noticias_copy.loc[index, 'Opinión de Agentes'])
    return noticias_copy


def analisis_1(doc, noticias):
    """
    Agrupa noticias por conjunto de provincias únicas y evita duplicaciones en el documento Word.
    """
    doc.add_heading('Pregunta 1 - Analisis Cuantitativo y Cualitativo', level=1)

    # Normalizar la columna de provincias para agrupar noticias con las mismas provincias
    noticias["Provincia_Normalizada"] = noticias["Provincia"].apply(
        lambda x: "-".join(sorted(set([p.strip() for p in x.split("-") if p.strip()])))
    )

    # Obtener combinaciones únicas de provincias
    provincias_unicas = noticias["Provincia_Normalizada"].unique()

    for provincia_group in sorted(provincias_unicas):  # Ordenar grupos de provincias
        doc.add_heading(provincia_group, level=2)  # Agregar título con las provincias agrupadas

        # Filtrar noticias que comparten exactamente el mismo grupo de provincias
        noticias_provincia = noticias[noticias["Provincia_Normalizada"] == provincia_group]

        for index, row in noticias_provincia.iterrows():
            doc.add_paragraph(f"({index + 1}) {row['Análisis Cuantitativo y Cualitativo']}", style="BodyText")

    return doc


def analisis_2_3_4(doc, noticias):
    """
    Agrega al documento Word los agentes territoriales agrupados por categoría,
    indicando el número de la noticia en la que aparecen.
    Ahora incluye datos de las Preguntas 2, 3 y 4.
    """
    doc.add_page_break()  # Salto de página para separar esta sección

    # Diccionario para almacenar los datos de cada pregunta
    preguntas = {
        "Pregunta 2 - Opinión de Agentes": "Opinión de Agentes",
        "Pregunta 3 - Análisis de Sentimientos": "Análisis de Sentimientos",
        "Pregunta 4 - Análisis Sentimientos Cuantitativo": "Análisis Sentimientos Cuantitativo"
    }

    for pregunta, columna in preguntas.items():
        doc.add_heading(pregunta, level=1)  # Título de la pregunta

        # Diccionario para agrupar agentes por categoría
        agentes_por_categoria = {}

        for index, row in noticias.iterrows():
            noticia_numero = index + 1  # Número de la noticia

            # Obtener la información de la columna correspondiente
            texto_columna = row[columna]

            if pd.notna(texto_columna):  # Verificar que no sea NaN
                lineas = texto_columna.split("\n")

                categoria_actual = None
                for linea in lineas:
                    if linea.startswith("*Agente"):  # Detecta una categoría de agentes
                        categoria_actual = linea.strip()
                        if categoria_actual not in agentes_por_categoria:
                            agentes_por_categoria[categoria_actual] = []

                    elif categoria_actual and "-" in linea:  # Es un agente dentro de la categoría
                        agente = linea.strip()
                        agentes_por_categoria[categoria_actual].append(f"({noticia_numero}) {agente}")

        # Agregar al documento Word las categorías y sus agentes
        for categoria, agentes in agentes_por_categoria.items():
            doc.add_heading(categoria, level=2)
            for agente in agentes:
                doc.add_paragraph(agente, style="BodyText")
    return doc


def analisis_estadistico_sentimientos_cuantitativos(noticias, doc, ruta_imagen): 
    # Lista de grupos (se deben incluir los dos puntos ':' al final para identificar el header)
    grupos = [
        "Agente Territorial Administraciones públicas-Local-Ayuntamiento:",
        "Agente Territorial Administraciones públicas-Nación:",
        "Agente Territorial Administraciones públicas-Regional Comunidad autónoma:",
        "Agente Territorial Administraciones públicas-Supramunicipal:",
        "Agente Territorial Ciudadanos:",
        "Agente Territorial Empresarios:",
        "Agente Territorial Grupos de presión:",
        "Agente Territorial Periodista:"
    ]
    
    # Diccionario para almacenar los valores numéricos de cada grupo
    datos = {grupo: [] for grupo in grupos}
    
    # Recorrer cada fila de la columna
    for texto in noticias['Análisis Sentimientos Cuantitativo']:
        # Para cada grupo, buscamos su bloque en el texto
        for grupo in grupos:
            # Debido a que en el texto el grupo podría venir rodeado de asteriscos (por ejemplo: "*Agente Territorial ...*:"), 
            # construimos un patrón que acepte opcionalmente dichos asteriscos.
            # Eliminamos el ':' final para formar el patrón base:
            grupo_sin_dos_puntos = grupo[:-1]
            # Patrón: opcionalmente un asterisco, el grupo, opcionalmente otro asterisco y luego dos puntos.
            patron_grupo = r"\*?" + re.escape(grupo_sin_dos_puntos) + r"\*?:"
            
            match = re.search(patron_grupo, texto)
            if match:
                # Una vez encontrado el grupo, tomamos el bloque de texto a partir del final del match
                inicio = match.end()
                fin = len(texto)
                # Para determinar el final del bloque, buscamos si aparece algún otro grupo después.
                for otro in grupos:
                    if otro == grupo:
                        continue
                    otro_sin_dos_puntos = otro[:-1]
                    patron_otro = r"\*?" + re.escape(otro_sin_dos_puntos) + r"\*?:"
                    match_otro = re.search(patron_otro, texto[inicio:])
                    if match_otro:
                        pos = inicio + match_otro.start()
                        if pos < fin:
                            fin = pos
                bloque = texto[inicio:fin]
                
                # En el bloque extraemos todos los números (aceptamos números negativos y decimales)
                numeros = re.findall(r"[-+]?\d*\.\d+|[-+]?\d+", bloque)
                for n in numeros:
                    try:
                        datos[grupo].append(float(n))
                    except Exception as e:
                        # En caso de error, se ignora ese valor
                        pass
                        
    # Calcular y escribir la media para cada grupo en 'doc'
    medias = {}
    desv_tip = {}
    doc.add_heading('Resultados estadísticos - Analisis de sentimientos cuantitativos', level=1)
    for grupo, valores in datos.items():
        if valores:
            media = np.mean(valores).round(2)
            medias[grupo] = media
            desviacionTipica = np.std(valores).round(2)
            desv_tip[grupo] = desviacionTipica
        else:
            medias[grupo] = None  # O se puede usar np.nan para indicar ausencia de datos
            desv_tip[grupo] = None
        # Formateamos la línea de salida
        doc.add_heading(grupo.strip(), level=2)
        salida = f"Media: {medias[grupo]}\nDesviación típica: {desv_tip[grupo]}\nNúmero de agentes: {len(valores)}"
        doc.add_paragraph(salida)
        imagen_path = f"{ruta_imagen}/{grupo.strip().replace(' ', '_').replace(':', '')}.png"
        create_grafico_barras_por_cada_valor(valores, imagen_path)
        doc.add_picture(imagen_path)
    
    return doc

def create_grafico_dispersion_datos(valores, imagen_path):
    plt.figure(figsize=(6, 2))  # Tamaño ajustado para un diseño horizontal
    plt.axhline(y=0, color='black', linestyle='-', linewidth=1.5, label='Línea base')  # Línea recta como guía
    plt.scatter(valores, [0] * len(valores), alpha=0.7, color='red', label='Valores')  # Colocar puntos sobre la línea
    plt.xlim(-1, 1)  # Eje fijo desde -1 a 1
    plt.ylim(-0.5, 0.5)  # Mantener el gráfico compacto verticalmente
    plt.title(f'Dispersión de los datos', fontsize=14)
    plt.xlabel('Puntuación sentimientos', fontsize=10)
    plt.yticks([])  # Ocultar etiquetas del eje y
    plt.grid(True, axis='x', linestyle='--', linewidth=0.5)  # Rejilla solo en el eje x
    plt.tight_layout()
    plt.savefig(imagen_path)
    plt.close()

def create_grafico_barras_por_cada_valor(valores, imagen_path):
    # Definir el rango deseado (por ejemplo, en pasos de 0.1)
    x_possible = np.linspace(-1, 1, 21)  # 21 valores: -1.0, -0.9, ..., 0.9, 1.0
    
    # Contar la frecuencia de cada valor en los datos
    counts = Counter(valores)
    
    # Para asegurar consistencia, redondeamos cada valor a 1 decimal
    y_values = [counts.get(round(x, 1), 0) for x in x_possible]
    
    # Crear el gráfico de barras
    plt.figure(figsize=(6, 3))
    plt.bar(x_possible, y_values, width=0.08, color='skyblue', edgecolor='black')
    plt.xlim(-1, 1)
    plt.title('Dispersión de los datos', fontsize=12)
    plt.xlabel('Puntuación sentimientos', fontsize=10)
    plt.ylabel('Nº de repeticiones', fontsize=10)
    
    # --- Sección para configurar el eje Y solo con las frecuencias que aparecen ---
    # Obtenemos las frecuencias únicas distintas de cero (si no deseas incluir el 0).
    # Si quieres incluir el 0 cuando aparezca, simplemente quita "if freq != 0".
    frecuencias_unicas = sorted({int(freq) for freq in y_values if freq != 0})
    
    # Ajustamos las 'yticks' para que solo muestre las frecuencias que existen
    plt.yticks(frecuencias_unicas, [str(f) for f in frecuencias_unicas])
    # ------------------------------------------------------------------------------

    plt.xticks(x_possible, rotation=45)
    plt.grid(True, axis='y', linestyle='--', linewidth=0.5)
    plt.tight_layout()
    plt.savefig(imagen_path)
    plt.close()

def crear_docx_completo(noticias, ruta_imagen):
    doc = Document()
    doc.add_heading('Análisis de noticias', level=0)
    doc = analisis_1(doc, noticias)
    doc = analisis_2_3_4(doc, noticias)
    doc = analisis_estadistico_sentimientos_cuantitativos(noticias, doc, ruta_imagen)
    return doc

def crear_docx_analisis_1(noticias):
    doc = Document()
    doc.add_heading('Análisis de noticias', level=0)
    doc = analisis_1(doc, noticias)
    return doc


def crear_docx_individual(noticias, ruta_carpeta):
    for index, row in noticias.iterrows():
        doc = Document()
        doc.add_heading(row['Titulo'], level=0)
        doc.add_paragraph(row['Contenido'])

        doc.add_heading('Analisis de sentimientos', level=1)
        doc.add_paragraph(row['Sentimiento explicado IA'])

        doc.add_heading('Pregunta 1 - Analisis Cuantitativo y Cualitativo', level=1)
        doc.add_paragraph(row['Análisis Cuantitativo y Cualitativo'])

        doc.add_heading('Pregunta 2 - Opinión de Agentes', level=1)
        doc.add_paragraph(row['Opinión de Agentes'])

        doc.add_heading('Pregunta 3 - Análisis de Sentimientos', level=1)
        doc.add_paragraph(row['Análisis de Sentimientos'])

        doc.add_heading('Pregunta 4 - Análisis de Sentimientos Cuantitativo', level=1)
        doc.add_paragraph(row['Análisis Sentimientos Cuantitativo'])

       

        if not os.path.exists(ruta_carpeta):
            os.makedirs(ruta_carpeta)  # Crea la carpeta si no existe
        ruta_archivo_docx = os.path.join(ruta_carpeta, f'Analisis_Noticia_{index+1}.docx')
        doc.save(ruta_archivo_docx)

def crear_docx_individual_analisis_1(noticias, ruta_carpeta):
    for index, row in noticias.iterrows():
        doc = Document()
        doc.add_heading(row['Titulo'], level=0)
        doc.add_paragraph(row['Contenido'])

        doc.add_heading('Analisis de sentimientos', level=1)
        doc.add_paragraph(row['Sentimiento explicado IA'])

        doc.add_heading('Pregunta 1 - Analisis Cuantitativo y Cualitativo', level=1)
        doc.add_paragraph(row['Análisis Cuantitativo y Cualitativo'])

        if not os.path.exists(ruta_carpeta):
            os.makedirs(ruta_carpeta)  # Crea la carpeta si no existe
        ruta_archivo_docx = os.path.join(ruta_carpeta, f'Analisis_Noticia_{index+1}.docx')
        doc.save(ruta_archivo_docx)